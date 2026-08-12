from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "ARGUS_FRONTEND_HANDOFF_JUANFER.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
CYAN = "00A9C7"
INK = "1D2939"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_CYAN = "E8F7FA"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
GREEN = "147D64"
GOLD = "A15C00"
RED = "9B1C1C"


def rgb(value):
    return RGBColor.from_string(value)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths_dxa = [int(width * 1440) for width in widths]
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=NAVY, size=font_size)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_properties.append(header_marker)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_para(doc, text="", style=None, size=11, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.10):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_para(doc, pieces, after=6, before=0, line=1.10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, options in pieces:
        run = p.add_run(text)
        set_run_font(run, **options)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: BLUE, 2: BLUE, 3: NAVY}
    before = {1: 16, 2: 12, 3: 8}
    after = {1: 8, 2: 6, 3: 4}
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after = Pt(after[level])
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=sizes[level], color=colors[level], bold=True)
    return p


def add_callout(doc, label, text, fill=LIGHT_CYAN, accent=CYAN):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.12
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color=accent, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F7F9FC")
    p_pr.append(shd)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=8.5, color=NAVY)
    return p


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, NAVY)):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("ARGUS  |  Documento de transferencia técnica")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("ARGUS  |  Uso interno  •  Página ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(footer)


def build():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "ARGUS - Guía de integración frontend para Juanfer"
    doc.core_properties.subject = "Arquitectura, cámaras, Tailscale y contratos de integración"
    doc.core_properties.author = "ARGUS"
    doc.core_properties.keywords = "ARGUS, frontend, cámaras, Tailscale, FastAPI, MongoDB"

    # Cover / masthead.
    add_para(doc, "ARGUS", size=12, color=CYAN, bold=True, after=8)
    add_para(doc, "Guía de integración frontend", size=28, color=NAVY, bold=True, after=4, line=1.0)
    add_para(doc, "Arquitectura, cámaras, Tailscale, estados y contrato para Juanfer", size=14, color=MUTED, after=22)
    add_table(
        doc,
        ["Documento", "Estado", "Referencia"],
        [["Transferencia técnica", "Implementación funcional + roadmap", "gustavo-backend / f5acf48"],
         ["Fecha", "12 de agosto de 2026", "ARGUS Core API + Edge"],
         ["Destinatario", "Juanfer / equipo frontend", "Preparado para integración web"]],
        [1.45, 2.7, 2.35],
        header_fill=LIGHT_BLUE,
        font_size=9.5,
    )
    add_callout(
        doc,
        "Idea central",
        "El frontend no procesa ni analiza el video. ARGUS Edge captura y analiza cada cámara; el Core API registra cámaras, eventos y estados; ARGUS Web presenta esa información en tiempo real.",
        fill=LIGHT_CYAN,
    )
    add_para(doc, "Documento interno de equipo", size=9.5, color=MUTED, italic=True, after=0)
    doc.add_page_break()

    add_heading(doc, "1. Qué es ARGUS y cuál es el objetivo", 1)
    add_para(doc, "ARGUS es un sistema distribuido de vigilancia y análisis visual. Su objetivo es recibir video desde varias cámaras, detectar actividad u objetos cerca de la fuente, registrar eventos y entregar al usuario una vista operativa clara: qué cámaras están conectadas, qué está ocurriendo y qué requiere atención.")
    add_para(doc, "La aplicación se diseña con tres capas:")
    add_bullet(doc, "Edge: proceso cercano a la cámara. Captura el video, ejecuta detección/tracking, dibuja overlays y emite heartbeats/eventos.")
    add_bullet(doc, "Core API: servicio central FastAPI. Registra cámaras, persiste eventos en MongoDB, calcula estados y publica actualizaciones WebSocket.")
    add_bullet(doc, "Frontend: dashboard web. Renderiza el grid de cámaras, timeline, alertas, filtros, estados y métricas.")
    add_callout(doc, "No confundir", "Un pod no es la cámara. La cámara es un dispositivo físico o URL; el pod/proceso Edge es el runtime que la representa y la procesa.", fill="FFF8E8", accent=GOLD)

    add_heading(doc, "2. Qué ya está hecho", 1)
    add_table(
        doc,
        ["Área", "Implementado", "Qué debe consumir Juanfer"],
        [
            ["API central", "FastAPI + MongoDB + configuración por entorno", "Base URL y respuestas JSON"],
            ["Salud", "GET /health valida realmente MongoDB", "Badge de API: healthy/degraded"],
            ["Cámaras", "CRUD, provisionamiento, heartbeat, status y FPS", "Grid, filtros y detalle de cámara"],
            ["Video", "MJPEG /stream por cada Edge", "<img src={stream_url}>"],
            ["Visión", "YOLO tracking, etiquetas, confianza, bbox e IDs", "Cajas/overlay ya vienen en el stream"],
            ["Tiempo real", "WebSocket /ws/events", "Timeline y alertas sin polling constante"],
            ["Rendimiento", "Captura separada de inferencia; latest-frame", "Mostrar FPS y latencia"],
            ["Multi-cámara", "Manifest JSON + launcher Edge", "Una tarjeta por cámara"],
            ["Red", "Compose separa core_net e iot_net; Tailscale planeado", "Usar IP Tailscale, nunca localhost remoto"],
        ],
        [1.15, 2.75, 2.6],
    )

    add_heading(doc, "3. Arquitectura y flujo de datos", 1)
    add_table(
        doc,
        ["Paso", "Componente", "Responsabilidad", "Salida"],
        [
            ["1", "Cámara / teléfono / IP", "Produce webcam, MJPEG o RTSP", "Frames"],
            ["2", "ARGUS Edge", "Captura, YOLO, tracking, HUD, reconexión", "MJPEG + heartbeat + eventos"],
            ["3", "Core API", "Valida, registra y publica estados", "JSON + WebSocket"],
            ["4", "MongoDB", "Persiste cámaras y eventos", "Historial"],
            ["5", "ARGUS Web", "Presenta video, estados y alertas", "Dashboard operativo"],
        ],
        [0.55, 1.45, 2.55, 1.95],
    )
    add_code(doc, "Cámara -> Edge (video + visión) -> Core API (estado/eventos) -> MongoDB\n                                      -> WebSocket -> Frontend\nFrontend -> GET /api/cameras -> stream_url de cada Edge")
    add_para(doc, "Regla importante: el video no se guarda en MongoDB. MongoDB guarda metadatos, estados y eventos. El stream permanece en el Edge y el frontend lo consume mediante la URL publicada por esa cámara.")

    add_heading(doc, "4. Contrato que debe implementar el frontend", 1)
    add_para(doc, "Base URL local: http://localhost:8000. En una demo remota, sustituir localhost por la IP Tailscale del nodo donde vive el Core API.")
    add_table(
        doc,
        ["Método", "Endpoint", "Uso frontend"],
        [
            ["GET", "/health", "Estado de API y base de datos"],
            ["GET", "/api/system/summary", "KPIs: cámaras totales, online, eventos del día, clientes realtime"],
            ["GET", "/api/cameras", "Carga inicial del grid; admite ?status=online y ?type=rtsp"],
            ["GET", "/api/cameras/{id}", "Detalle y estado actual de una cámara"],
            ["POST", "/api/cameras/configure", "Provisionamiento administrativo; devuelve edge_config y start_command"],
            ["GET", "/api/events", "Timeline; admite camera_id, type, status y limit"],
            ["PATCH", "/api/events/{id}", "Acknowledge/resolved de una alerta"],
            ["WS", "/ws/events", "Actualizaciones instantáneas de eventos y cámaras"],
        ],
        [0.65, 2.45, 3.4],
    )
    add_heading(doc, "4.1 Modelo Camera", 2)
    add_code(doc, '''{
  "id": "CAM-001",
  "name": "Camera Laptop Gustavo",
  "type": "integrated",
  "stream_url": "http://100.x.y.z:8091/stream",
  "source": "0",
  "vision_mode": "cctv",
  "status": "online",
  "fps": 10.4,
  "location": "Laboratorio",
  "metadata": {"network": {"iot_segment": "iot-cameras"}},
  "last_heartbeat": "2026-08-12T04:24:29Z"
}''')
    add_heading(doc, "4.2 Modelo Event", 2)
    add_code(doc, '''{
  "id": "EVT-ABC123",
  "camera_id": "CAM-001",
  "type": "object_detected",
  "object_name": "person",
  "confidence": 0.87,
  "status": "new",
  "metadata": {"track_id": 4, "bbox": [120, 80, 360, 500]},
  "timestamp": "2026-08-12T04:24:29Z"
}''')
    add_callout(doc, "Video", "El stream actual es MJPEG. En React, Vue o HTML se puede renderizar directamente con <img src={camera.stream_url} />. No usar <video> para este endpoint MJPEG.", fill=LIGHT_CYAN, accent=CYAN)

    add_heading(doc, "5. WebSocket y estados en tiempo real", 1)
    add_para(doc, "Conectar una sola vez a ws://localhost:8000/ws/events y actualizar el estado local del dashboard según el campo type. El socket recibe eventos creados/actualizados/eliminados y cambios de estado de cámaras.")
    add_table(
        doc,
        ["type", "Contenido", "Acción UI"],
        [
            ["event_created", "event", "Agregar al timeline, contador y alerta"],
            ["event_updated", "event", "Actualizar status y detalle"],
            ["event_deleted", "event_id", "Retirar del timeline si corresponde"],
            ["camera_status", "camera + event", "Actualizar badge, color y disponibilidad"],
        ],
        [1.45, 2.25, 2.8],
    )
    add_heading(doc, "5.1 Máquina de estados de cámara", 2)
    add_table(
        doc,
        ["Estado", "Significado", "Visual recomendado"],
        [
            ["online", "Heartbeat reciente y captura activa", "Verde + stream disponible"],
            ["degraded", "El agente existe, pero captura o dependencia presenta problemas", "Ámbar + warning"],
            ["offline", "Sin heartbeat dentro del timeout o cámara sin registro activo", "Rojo/gris + reconectar"],
        ],
        [1.1, 3.6, 1.8],
    )
    add_para(doc, "El Edge envía heartbeat aproximadamente cada cinco segundos. El Core API considera offline una cámara cuyo último heartbeat supera CAMERA_HEARTBEAT_TIMEOUT_SECONDS, actualmente 15 segundos. El frontend no debe asumir que stream_url accesible significa cámara saludable: debe combinar stream y status.")

    add_heading(doc, "6. Tailscale: cómo conectar celulares, laptops y cámaras", 1)
    add_para(doc, "Tailscale será la red privada sobre la que los nodos Edge, el Core API y los equipos del frontend se encuentran por IP estable de tailnet. La cámara física no necesita estar dentro de Kubernetes: necesita que el host que ejecuta Edge sea accesible por Tailscale.")
    add_heading(doc, "6.1 Topología propuesta", 2)
    add_table(
        doc,
        ["Nodo", "Instala Tailscale", "Expone / consume"],
        [
            ["Gustavo / Core", "Sí", "API :8000, Mongo solo interno, opcional frontend"],
            ["Laptop con webcam", "Sí", "Edge :8091 y source 0"],
            ["Celular con cámara", "Sí, si la app de cámara permite acceso por tailnet", "Edge o stream IP :8092"],
            ["Juanfer / frontend", "Sí para demo remota", "Core :8000 y streams Edge :809x"],
        ],
        [1.55, 2.05, 2.9],
    )
    add_heading(doc, "6.2 Instalación y prueba por nodo", 2)
    add_number(doc, "Instalar Tailscale en Windows, Android/iOS o Linux e iniciar sesión en el mismo tailnet.")
    add_number(doc, "En cada host ejecutar tailscale ip -4 y anotar la IP 100.x.y.z.")
    add_number(doc, "Desde otro nodo validar conectividad con tailscale ping <nombre-o-ip>.")
    add_number(doc, "Probar el servicio usando la IP Tailscale: http://100.x.y.z:8091/health y /stream.")
    add_number(doc, "Provisionar la cámara con edge_host igual a esa IP, no localhost.")
    add_code(doc, "tailscale ip -4\ntailscale status\ntailscale ping <edge-host>\n\n# Ejemplo de URL remota\nhttp://100.101.102.103:8091/stream")
    add_callout(doc, "0.0.0.0", "Es una dirección de escucha del servidor, no una dirección navegable. El navegador debe usar localhost en el mismo host o la IP Tailscale/LAN del host remoto.", fill="FFF8E8", accent=GOLD)
    add_heading(doc, "6.3 Segmentación IoT", 2)
    add_para(doc, "iot_segment es actualmente una etiqueta lógica que viaja en metadata.network. Compose separa argus_core_net (MongoDB + API) de argus_iot_net (API + Edge), por lo que MongoDB no queda dentro del segmento IoT. La segmentación física se completa con Tailscale grants/ACLs, firewall y puertos explícitos.")
    add_bullet(doc, "Permitir frontend -> Core API en 8000.")
    add_bullet(doc, "Permitir frontend -> Edge únicamente en los puertos de stream necesarios, por ejemplo 8091-8099.")
    add_bullet(doc, "Permitir Edge -> Core API en 8000 para registro, heartbeat y eventos.")
    add_bullet(doc, "No exponer MongoDB al frontend ni a cámaras.")
    add_para(doc, "Para el tailnet conviene definir tags como tag:argus-core, tag:argus-edge y tag:argus-web, y otorgar solo los accesos requeridos. Tailscale recomienda grants para nuevas políticas; Tailscale Serve puede publicar servicios solo dentro del tailnet. Funnel debe reservarse para una necesidad explícita de exposición pública.", size=10.5)

    add_heading(doc, "7. Varias cámaras y la idea de pods virtuales", 1)
    add_para(doc, "Cada cámara se registra como un recurso independiente en MongoDB y tiene un Edge independiente. En local, cada Edge es un proceso PowerShell/uvicorn; en Linux puede ser un contenedor; en Kubernetes será un Pod o Deployment por fuente RTSP/MJPEG. La identidad lógica no cambia: camera_id, stream_url, status, fps, metadata y last_heartbeat.")
    add_table(
        doc,
        ["Cámara", "Proceso Edge", "Puerto", "Fuente", "Recurso visible"],
        [
            ["CAM-001", "Edge laptop", "8091", "0", "Camera Laptop Gustavo"],
            ["CAM-002", "Edge IP/phone", "8092", "rtsp://...", "Camera IP Laboratorio"],
            ["CAM-003", "Edge USB/host B", "8093", "1", "Camera Oficina"],
        ],
        [0.85, 1.55, 0.75, 1.65, 1.7],
    )
    add_heading(doc, "7.1 Arranque en Windows", 2)
    add_para(doc, "El repositorio ya incluye scripts/cameras.example.json y scripts/start-edge-fleet.ps1. Copiar el ejemplo a cameras.local.json, activar las entradas necesarias y ejecutar el launcher. Cada proceso se registra automáticamente en el Core API.")
    add_code(doc, "Copy-Item .\\scripts\\cameras.example.json .\\scripts\\cameras.local.json\n# Editar cameras.local.json: enabled=true, port único y source real\n.\\scripts\\start-edge-fleet.ps1\n\nInvoke-RestMethod http://localhost:8000/api/cameras")
    add_callout(doc, "Escalamiento futuro", "En Kubernetes no se debe lanzar un pod genérico que intente abrir todas las cámaras. Lo más estable es un Edge por stream o un deployment por grupo controlado, con ConfigMap/Secret y un camera_id inmutable.", fill=LIGHT_CYAN, accent=CYAN)

    add_heading(doc, "8. Diseño recomendado para ARGUS Web", 1)
    add_para(doc, "La primera versión del frontend debe priorizar observabilidad y operación. La inteligencia ya llega desde Edge como cajas, etiquetas y eventos; el frontend debe hacer que esa información sea fácil de leer y accionar.")
    add_table(
        doc,
        ["Componente", "Responsabilidad", "Fuente"],
        [
            ["CameraGrid", "Una tarjeta por cámara, responsive", "GET /api/cameras"],
            ["CameraCard", "Stream, nombre, estado, FPS, modo, location", "Camera + stream_url"],
            ["AlertRail", "Alertas nuevas y acknowledge/resolved", "WS + /api/events"],
            ["Timeline", "Eventos filtrados por cámara/tipo/estado", "GET /api/events"],
            ["SystemSummary", "Cámaras totales/online y eventos de hoy", "GET /api/system/summary"],
            ["CameraDetail", "Detalle, latencia, labels y configuración", "GET /api/cameras/{id}"],
            ["ConnectionStatus", "API, DB, WebSocket y Edge", "health + camera status"],
        ],
        [1.45, 3.35, 1.6],
    )
    add_heading(doc, "8.1 Comportamiento esperado", 2)
    add_bullet(doc, "Al cargar: pedir /health, /api/system/summary y /api/cameras en paralelo.")
    add_bullet(doc, "Crear el grid con skeleton cards mientras llegan los datos.")
    add_bullet(doc, "Usar stream_url directamente en una etiqueta img y mostrar placeholder si status no es online.")
    add_bullet(doc, "Abrir WebSocket después de la carga inicial; aplicar eventos por ID para no recargar toda la pantalla.")
    add_bullet(doc, "Si el socket se desconecta, mostrar Reconnecting y reintentar con backoff; no duplicar listeners.")
    add_bullet(doc, "No mostrar una alerta de persona/objeto cada frame: el Edge ya aplica cooldown de eventos.")
    add_heading(doc, "8.2 Texto para enviarle a Juanfer", 2)
    add_callout(doc, "Mensaje", "Juanfer, ARGUS Web debe consumir el Core API en /api/cameras, /api/events y /api/system/summary, y escuchar /ws/events. Cada cámara trae su propio stream_url MJPEG: renderízalo con <img>. El backend ya entrega status, fps, vision_mode, detections, object_name, confidence, bbox y track_id; el frontend debe concentrarse en grid, estados, timeline, filtros y alertas en tiempo real.", fill=LIGHT_CYAN, accent=CYAN)

    add_heading(doc, "9. Despliegue y pruebas de integración", 1)
    add_heading(doc, "9.1 Core local", 2)
    add_code(doc, "Copy-Item .env.example .env\ndocker compose up -d --build mongodb api\ndocker compose ps\nInvoke-RestMethod http://localhost:8000/health")
    add_heading(doc, "9.2 Edge de la laptop", 2)
    add_code(doc, ".\\scripts\\start-edge.ps1 `\n  -CameraId CAM-001 `\n  -CameraName \"Camera Laptop Gustavo\" `\n  -CameraType integrated `\n  -CameraSource 0 `\n  -Port 8091 `\n  -Vision `\n  -VisionMode cctv")
    add_heading(doc, "9.3 Checklist de aceptación", 2)
    add_bullet(doc, "GET /health del Core devuelve database=connected.")
    add_bullet(doc, "GET /health del Edge devuelve status=online, detector cargado y detection_error=null.")
    add_bullet(doc, "GET /api/cameras contiene stream_url accesible desde el navegador del frontend.")
    add_bullet(doc, "El stream muestra video y cajas de detección sin congelar la captura.")
    add_bullet(doc, "Al aparecer una persona/objeto se registra object_detected y llega event_created por WebSocket.")
    add_bullet(doc, "Al detener Edge, la cámara pasa a offline después del timeout.")
    add_bullet(doc, "Con dos cámaras, cada una conserva su video, estado, FPS y eventos por separado.")

    add_heading(doc, "10. Límites actuales y roadmap", 1)
    add_table(
        doc,
        ["Ahora", "Siguiente mejora", "Resultado esperado"],
        [
            ["MJPEG por cámara", "Gateway WebRTC/RTSP", "Menor latencia y mejor escala de viewers"],
            ["YOLO preentrenado", "Dataset propio + fine-tuning", "Mejor detección en habitaciones reales"],
            ["Cajas y tracking", "Zonas y reglas", "Alertar solo dentro de áreas relevantes"],
            ["Eventos básicos", "Snapshots y retención", "Evidencia visual asociada al evento"],
            ["WebSocket in-process", "Redis/pub-sub", "Eventos consistentes con varias réplicas del API"],
            ["Sin auth en API", "JWT/Keycloak + roles", "Acceso seguro por usuario/dispositivo"],
            ["Modos preparados", "Pose/actividad y expresión", "Análisis temporal separado del detector de objetos"],
        ],
        [1.65, 2.65, 2.1],
    )
    add_callout(doc, "Prioridad recomendada", "Primero integrar el frontend con cámaras, estados y eventos. Después medir rendimiento con dos cámaras. Luego entrenar clases propias con datos reales. Finalmente añadir zonas, snapshots, autenticación y análisis de actividad.", fill="FFF8E8", accent=GOLD)

    add_heading(doc, "11. Referencias técnicas", 1)
    add_para(doc, "Referencias oficiales usadas para orientar la arquitectura y las decisiones de red:", after=4)
    add_bullet(doc, "Ultralytics Tracking: https://docs.ultralytics.com/modes/track")
    add_bullet(doc, "Ultralytics Training: https://docs.ultralytics.com/modes/train")
    add_bullet(doc, "Ultralytics Benchmark/Export: https://docs.ultralytics.com/modes/benchmark")
    add_bullet(doc, "OpenCV VideoCapture properties: https://docs.opencv.org/4.9.0/d4/d15/group__videoio__flags__base.html")
    add_bullet(doc, "Tailscale Windows: https://tailscale.com/docs/install/windows")
    add_bullet(doc, "Tailscale grants: https://tailscale.com/docs/features/access-control/grants")
    add_bullet(doc, "Tailscale Serve: https://tailscale.com/docs/reference/tailscale-cli/serve")

    add_heading(doc, "12. Resumen final para el equipo", 1)
    add_para(doc, "ARGUS ya tiene la base de un sistema distribuido de vigilancia: cada cámara puede vivir en un Edge distinto, el Core API centraliza estado y eventos, MongoDB conserva la historia y el frontend se limita a convertir esos datos en una experiencia operativa. La decisión arquitectónica clave es mantener separadas captura, inferencia, persistencia y presentación. Esa separación permite empezar con la webcam de la laptop, sumar celulares/IP cameras y luego mover los Edge a contenedores o pods sin cambiar el contrato que consume Juanfer.")
    add_callout(doc, "Entrega", "Documento preparado sobre la rama gustavo-backend y el estado funcional publicado hasta f5acf48.", fill=LIGHT_CYAN, accent=CYAN)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
